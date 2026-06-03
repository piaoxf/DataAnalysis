from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent


answers = {
    "a-1": (
        "No. Regressing an outcome on as many variables as possible and then reporting only "
        "statistically significant coefficients is likely to produce misleading policies. With many "
        "candidate variables, some will look significant by chance, and the selected coefficients are "
        "biased upward because they were chosen after seeing the data. Statistical significance also "
        "does not establish causality, practical importance, or robustness. A policy recommendation "
        "should be based on a pre-specified question, credible identification strategy, effect sizes, "
        "uncertainty, and validation."
    ),
    "a-2": (
        "More data helps with sampling noise, but it does not by itself make the proposed procedure "
        "reliable. Consider a simple experiment: simulate an educational outcome and 1,000 unrelated "
        "covariates, all with true effect zero, then test each coefficient at the 5% level. Even with "
        "perfectly generated data, about 50 variables are expected to be declared significant just by "
        "chance if no multiple-testing correction is used. If the number of searched variables grows "
        "with the sample size, false discoveries can remain a serious problem. In observational data, "
        "larger samples can also make tiny, unimportant associations statistically significant, while "
        "confounding and measurement bias remain. If the model were fixed in advance, correctly "
        "specified, and all assumptions held, increasing sample size would help distinguish true "
        "nonzero effects from zero effects. But the proposed workflow is a data-mining procedure: it "
        "uses the same data to search, select, and claim discoveries. To find true policy effects, one "
        "should pre-specify hypotheses, use randomized or quasi-experimental designs where possible, "
        "adjust for multiple comparisons, report all analyses, and validate findings on new data."
    ),
    "b-2": (
        "No. A significant country-level correlation between chocolate consumption and Nobel laureates "
        "does not show a relationship between chocolate and intelligence. It may reflect confounding "
        "variables such as national income, education, research investment, population size, or reporting "
        "differences. It is also an ecological comparison: relationships between countries need not hold "
        "for individuals. The result may be useful for generating a hypothesis, but it is not evidence "
        "that chocolate improves intelligence."
    ),
    "b-3": (
        "They should study the question with a design that directly measures individuals and reduces "
        "confounding. A better observational study would collect individual-level chocolate consumption, "
        "baseline cognitive measures, age, education, income, health, family background, and other diet "
        "variables, then pre-specify the model and outcomes. A stronger design would randomly assign "
        "participants, when ethical and feasible, to controlled diets that differ in chocolate or cocoa "
        "content while keeping calories and other nutrients similar. They should report effect sizes, "
        "confidence intervals, all tested outcomes, and replication evidence."
    ),
    "b-4": (
        "They have some evidence that the assigned diet changed maze-solving time in this experiment, "
        "because randomization helps support a causal interpretation for the measured outcome in these "
        "mice. However, p < 0.05 alone is not enough to claim broadly that chocolate improves cognitive "
        "power. They should examine the effect size, confidence interval, experimental quality, whether "
        "the maze outcome was pre-specified, and whether multiple outcomes were tested. The conclusion "
        "should also be limited to mice under this specific diet and should be replicated before making "
        "a general scientific claim."
    ),
    "b-5": (
        "No. This is selective reporting, or p-hacking. If the lab tested about 100 features and then "
        "reported only those with p-values below 0.05, some significant results are expected by chance. "
        "They should disclose the full set of tests, distinguish planned from exploratory analyses, and "
        "adjust for multiple comparisons or validate the findings in new data."
    ),
    "c": (
        "The first title is too strong if it is based only on p < 0.05. A small p-value indicates that "
        "the observed data would be unusual under the null model, but it does not measure the size or "
        "clinical importance of the treatment effect. A better title would mention evidence of an effect "
        "and report the estimated effect size. The second title is clearly inappropriate: p < 0.05 does "
        "not mean the drug has over a 95% success rate, nor that the hypothesis is 95% likely to be true."
    ),
    "d": (
        "No. Failing to reject the null does not prove that HR spending has no effect. With only 25 years "
        "of data, the regression may have low power, and the estimate is described as large even though "
        "uncertain. The result should be interpreted as inconclusive, not as evidence of no benefit. The "
        "boss should consider the confidence interval, practical effect size, possible confounding, and "
        "whether the model is appropriate for a budget decision."
    ),
    "e": (
        "False. Successful replication strengthens a scientific claim, especially when the design is "
        "pre-specified, well powered, and low bias. However, replicated statistical significance is still "
        "not absolute proof. The claim should be framed with effect sizes, uncertainty, study design, "
        "model assumptions, and the scope of the population studied."
    ),
    "f": (
        "No. Reporting only statistically significant tests hides the number of tests that were tried and "
        "inflates the apparent strength of the evidence. Readers cannot judge the false-positive risk, "
        "and the paper becomes biased toward exciting results. He should report all pre-specified tests, "
        "label exploratory analyses clearly, and use appropriate multiple-testing corrections or "
        "independent validation."
    ),
    "g": (
        "True. A p-value is calculated under a specified statistical model. If the model does not match "
        "reality, a significant p-value may appear even when the substantive null hypothesis is true. "
        "Examples include nonrandom sampling, dependence among observations, measurement bias, omitted "
        "confounders, optional stopping, or using the wrong error distribution. Thus a small p-value is "
        "evidence against a model, not automatic proof that the scientific null is false."
    ),
}


def set_margins(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if name == "Heading 2" else 16)
        style.paragraph_format.space_after = Pt(6)


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("555555")


def add_answer(doc, label, prompt_short, answer):
    doc.add_heading(f"({label}) {prompt_short}", level=2)
    doc.add_paragraph(answer)


def build_submission_report():
    doc = Document()
    set_margins(doc)
    configure_styles(doc)

    add_title(
        doc,
        "Written Report - Problem 1.3",
        "Interpreting p-values, statistical significance, and scientific evidence",
    )
    doc.add_paragraph("Name: ____________________")
    doc.add_paragraph(
        "Note: Part (b-1) is omitted from this written report because the assignment states that it is autograded."
    )

    doc.add_heading("Problem 1.3 Answers", level=1)
    add_answer(doc, "a-1", "Variable selection for education policy", answers["a-1"])
    add_answer(doc, "a-2", "Will more data fix the procedure?", answers["a-2"])
    add_answer(doc, "b-2", "Chocolate consumption and intelligence", answers["b-2"])
    add_answer(doc, "b-3", "A better study design", answers["b-3"])
    add_answer(doc, "b-4", "Randomized mouse experiment", answers["b-4"])
    add_answer(doc, "b-5", "Selective reporting after many tests", answers["b-5"])
    add_answer(doc, "c", "Interpreting a drug trial headline", answers["c"])
    add_answer(doc, "d", "Non-significant HR spending estimate", answers["d"])
    add_answer(doc, "e", "Replication and scientific claims", answers["e"])
    add_answer(doc, "f", "Reporting only significant tests", answers["f"])
    add_answer(doc, "g", "Significance under model mismatch", answers["g"])

    doc.add_heading("References", level=1)
    refs = [
        "Wasserstein, R. L., & Lazar, N. A. (2016). The ASA's Statement on p-Values: Context, Process, and Purpose. The American Statistician, 70(2), 129-133. https://doi.org/10.1080/00031305.2016.1154108",
        "Ioannidis, J. P. A. (2005). Why Most Published Research Findings Are False. PLOS Medicine, 2(8), e124. https://doi.org/10.1371/journal.pmed.0020124",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    out = OUT_DIR / "Problem_1_3_Report_Submission.docx"
    doc.save(out)
    return out


def build_explanation_doc():
    doc = Document()
    set_margins(doc)
    configure_styles(doc)

    add_title(
        doc,
        "Problem 1.3 中文详细说明",
        "p-value、显著性、选择性汇报与论文阅读指南",
    )

    doc.add_heading("一、这道题的总线索", level=1)
    doc.add_paragraph(
        "这道题不是在考复杂计算，而是在考如何正确解释 p-value。核心原则是："
        "p-value 只是在某个统计模型和零假设成立时，观察到当前或更极端数据的概率。"
        "它不是“假设为真的概率”，不是“结果由随机产生的概率”，也不是“效应大小”。"
    )
    doc.add_paragraph(
        "所以每个小问都围绕同一个判断：一个显著结果是否能支持因果、重要性、政策、"
        "科学发现或新闻标题？通常答案是：不能只靠 p < 0.05，还要看研究设计、先验合理性、"
        "多重检验、选择性报告、效应大小、置信区间和可重复性。"
    )

    doc.add_heading("二、逐题思路", level=1)
    explanations = [
        ("a-1", "把尽可能多的变量放进回归，再只挑显著变量，是典型的数据挖掘。变量越多，偶然显著越多；而且观察性回归不能自动给出因果。"),
        ("a-2", "更多数据可以降低随机误差，但不能修复错误研究流程。如果模型不预先设定、变量很多、报告有选择性，甚至大样本也会把微小或偏倚导致的关系变成显著。"),
        ("b-1", "自动评分题应选 No：巧克力和 Nobel 奖显著相关并不等于有真实因果关系。"),
        ("b-2", "国家层面的相关不能推出个人层面的关系，这是 ecological fallacy；GDP、教育、科研投入、人口规模等都可能混杂。"),
        ("b-3", "改进方法是做个体层面研究，最好预注册、明确变量和结果；若伦理和可行性允许，用随机实验更能回答因果问题。"),
        ("b-4", "小鼠随机实验比国家相关强，但 p < 0.05 只能说明有一些证据；还要看效应大小、是否多重比较、是否可重复、能否推广到“认知能力”。"),
        ("b-5", "先检验 100 个变量，再只写显著的几个，是选择性汇报，会夸大证据强度。"),
        ("c", "p < 0.05 不等于强效，也不等于 95% 成功率。新闻标题必须区分统计显著、效应大小和治疗成功概率。"),
        ("d", "不显著不代表没有效应。25 年样本很小，可能只是检验功效不足；大效应但不显著说明不确定性大。"),
        ("e", "这句话是 False。重复实验显著会增强证据，但仍不是绝对证明；科学结论还要看设计、偏倚、效应大小和适用范围。"),
        ("f", "只报告显著结果不可以，因为读者看不到失败的检验，也无法判断 false positive 风险。"),
        ("g", "True。p-value 依赖模型；如果独立性、抽样、分布、测量或混杂假设错了，小 p-value 可能只是模型错了。"),
    ]
    for label, text in explanations:
        doc.add_heading(f"({label})", level=2)
        doc.add_paragraph(text)

    doc.add_heading("三、Ioannidis 论文怎么读", level=1)
    steps = [
        "先读摘要和 Summary：抓住结论不是“所有论文都假”，而是很多研究领域中，发表的显著发现有很低的阳性预测值。",
        "再读 Modeling the Framework：只需要理解 R、power、alpha、PPV。R 是研究前真关系和假关系的比例；power 是真效应被发现的概率；alpha 是假阳性率。",
        "接着读 Bias 和 Testing by Several Independent Teams：这两节解释为什么灵活分析、选择性汇报、多个团队反复寻找显著结果，会让发表结果更不可靠。",
        "重点读 Corollaries：小样本、小效应、海量假设、灵活分析、利益冲突和热门领域，都会降低显著发现为真的概率。",
        "最后读 How Can We Improve the Situation：作者主张提高功效、预注册、减少偏倚、看整体证据，而不是追逐单个 p < 0.05。",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("四、Ioannidis 的核心思想", level=1)
    doc.add_paragraph(
        "Ioannidis 的核心思想可以用一句话概括：一个研究结果显著，并不表示它很可能是真的；"
        "它为真的概率取决于研究前这个假设有多可信、研究功效多高、偏倚多大、检验了多少关系、"
        "以及是否存在选择性汇报。"
    )
    doc.add_paragraph(
        "这可以类比医学检测：如果一种疾病本来极其罕见，即使检测阳性，也可能是假阳性。"
        "科研也是一样。如果在成千上万个低先验概率的假设里找 p < 0.05，再只发表显著结果，"
        "那么发表出来的“发现”中会混入大量假阳性。"
    )

    doc.add_heading("五、和 ASA p-value 声明的关系", level=1)
    doc.add_paragraph(
        "ASA 声明强调六个常见误区：p-value 只能衡量数据和某个模型的不相容程度；"
        "它不能告诉我们零假设为真的概率；科学结论不能只看是否过 0.05；"
        "正确推断需要完整报告和透明度；p-value 不是效应大小或实际重要性的度量；"
        "单独一个 p-value 不是证据强弱的完整指标。"
    )
    doc.add_paragraph(
        "因此，这道题几乎每个小问都可以用同一句话检查：这个结论是不是把“统计显著”"
        "误读成了“真实、因果、重要、可推广或已经证明”？如果是，就要谨慎。"
    )

    doc.add_heading("参考文献", level=1)
    doc.add_paragraph(
        "Wasserstein, R. L., & Lazar, N. A. (2016). The ASA's Statement on p-Values: Context, Process, and Purpose. The American Statistician, 70(2), 129-133."
    )
    doc.add_paragraph(
        "Ioannidis, J. P. A. (2005). Why Most Published Research Findings Are False. PLOS Medicine, 2(8), e124."
    )

    out = OUT_DIR / "Problem_1_3_Detailed_Explanation_CN.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    for path in [build_submission_report(), build_explanation_doc()]:
        print(path)
